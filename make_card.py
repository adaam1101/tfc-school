#!/usr/bin/env python3
"""Generate a fillable, ID-card-sized Word template for TFC.

Design brief (from the user):
  - Standard ID-card size (CR80: 3.375 in x 2.125 in)
  - White card
  - School logo sitting in the background as a faint transparent watermark
  - Field labels (Name / Date of birth / ...) in the SAME blue as the logo
"""
import copy
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----- Brand -----------------------------------------------------------------
LOGO_BLUE = RGBColor(0x04, 0x43, 0x6E)          # sampled from the logo
LOGO_BLUE_HEX = "04436E"
GREY_TEXT = RGBColor(0x33, 0x3B, 0x45)
LOGO_SRC = "client/public/tfc-logo.png"

CARD_W_IN = 3.375
CARD_H_IN = 2.125
DPI = 300
CARD_W_PX = int(CARD_W_IN * DPI)   # 1012
CARD_H_PX = int(CARD_H_IN * DPI)   # 637

BG_PATH = "card_bg.png"


# ----- 1. Build the card background PNG --------------------------------------
def build_background():
    """White card + faint centered logo watermark + blue accents."""
    W, H = CARD_W_PX, CARD_H_PX
    card = Image.new("RGB", (W, H), "white")

    # --- faint logo watermark, centered, blended toward white ---
    logo = Image.open(LOGO_SRC).convert("RGB")
    # scale logo to ~78% of card height
    target = int(H * 0.92)
    logo = logo.resize((target, target), Image.LANCZOS)
    # blend each pixel toward white -> very light watermark
    alpha = 0.10
    wm = Image.blend(Image.new("RGB", logo.size, "white"), logo, alpha)
    lx = (W - target) // 2
    ly = (H - target) // 2
    card.paste(wm, (lx, ly))

    from PIL import ImageDraw
    d = ImageDraw.Draw(card)

    # --- top + bottom accent bars in logo blue ---
    blue = (0x04, 0x43, 0x6E)
    sky = (0x2E, 0x6E, 0xA8)
    bar = int(H * 0.055)
    d.rectangle([0, 0, W, bar], fill=blue)
    d.rectangle([0, H - bar, W, H], fill=blue)
    # thin sky accent under the top bar
    d.rectangle([0, bar, W, bar + max(2, int(H * 0.012))], fill=sky)
    d.rectangle([0, H - bar - max(2, int(H * 0.012)), W, H - bar], fill=sky)

    # --- subtle inner border ---
    inset = int(H * 0.085)
    d.rectangle([inset, inset, W - inset, H - inset],
                outline=(0xCF, 0xDD, 0xEA), width=2)

    card.save(BG_PATH, dpi=(DPI, DPI))


# ----- helpers for the docx --------------------------------------------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def no_cell_margins_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tcPr.append(borders)


def add_floating_bg(paragraph, img_path, width_in, height_in):
    """Add an image anchored behindDoc=1 (true 'behind text' watermark)."""
    run = paragraph.add_run()
    # add as an inline first, then convert to a floating anchor
    inline_w = Emu(int(width_in * 914400))
    inline_h = Emu(int(height_in * 914400))
    run.add_picture(img_path, width=inline_w, height=inline_h)

    # grab the inline element and rebuild as <wp:anchor behindDoc="1">
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    anchor = OxmlElement("wp:anchor")
    for k, v in inline.attrib.items():
        anchor.set(k, v)
    anchor.set("behindDoc", "1")
    anchor.set("simplePos", "0")
    anchor.set("locked", "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "1")
    anchor.set("relativeHeight", "0")

    # simplePos
    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0"); sp.set("y", "0")
    anchor.append(sp)

    # horizontal position: aligned to column, offset 0
    posH = OxmlElement("wp:positionH")
    posH.set("relativeFrom", "column")
    ph = OxmlElement("wp:posOffset"); ph.text = "0"
    posH.append(ph)
    anchor.append(posH)

    posV = OxmlElement("wp:positionV")
    posV.set("relativeFrom", "paragraph")
    pv = OxmlElement("wp:posOffset"); pv.text = "0"
    posV.append(pv)
    anchor.append(posV)

    # carry over the inline children (extent, docPr, graphic, etc.)
    for child in list(inline):
        anchor.append(copy.deepcopy(child))

    # behind text -> wrapNone
    wrapNone = OxmlElement("wp:wrapNone")
    # insert wrapNone after positionV, before docPr
    docpr = anchor.find(qn("wp:docPr"))
    anchor.insert(list(anchor).index(docpr), wrapNone)

    drawing.remove(inline)
    drawing.append(anchor)


def field(cell, label, value=""):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(label.upper() + "   ")
    r.bold = True
    r.font.size = Pt(6.5)
    r.font.color.rgb = LOGO_BLUE
    r.font.name = "Arial"
    v = p.add_run(value if value else "______________________")
    v.font.size = Pt(8)
    v.font.color.rgb = GREY_TEXT if value else RGBColor(0xB6, 0xC2, 0xCE)
    v.font.name = "Arial"
    return p


# ----- 2. Build the Word document --------------------------------------------
def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("TFC — Official Identity Card  (print at 100%, 3.375\" × 2.125\")")
    tr.font.size = Pt(9)
    tr.font.color.rgb = RGBColor(0x8A, 0x97, 0xA5)
    tr.italic = True

    # single-cell table sized exactly to the card
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(CARD_W_IN)
    # lock row height to card height
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(int(CARD_H_IN * 1440)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

    no_cell_margins_borders(cell)

    # the cell's first paragraph hosts the floating background image
    bgp = cell.paragraphs[0]
    bgp.paragraph_format.space_after = Pt(0)
    bgp.paragraph_format.space_before = Pt(0)
    add_floating_bg(bgp, BG_PATH, CARD_W_IN, CARD_H_IN)

    # header line (school name) — small top padding to clear the blue bar
    pad = cell.add_paragraph()
    pad.paragraph_format.space_after = Pt(0)
    pad.add_run().font.size = Pt(3)

    head = cell.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_after = Pt(0)
    hr = head.add_run("TRAINING FORMATION CENTER")
    hr.bold = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = LOGO_BLUE
    hr.font.name = "Arial"

    sub = cell.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    sr = sub.add_run("STUDENT IDENTITY CARD")
    sr.font.size = Pt(6.5)
    sr.font.color.rgb = RGBColor(0x2E, 0x6E, 0xA8)
    sr.font.name = "Arial"

    # fields
    field(cell, "Name")
    field(cell, "Date of birth")
    field(cell, "Course")
    field(cell, "Card ID")
    field(cell, "Valid thru")
    field(cell, "Signature")

    doc.save("TFC-ID-Card-Template.docx")
    print("saved TFC-ID-Card-Template.docx")


# ----- 3. Whole-card visual preview (for the user) ---------------------------
def build_preview():
    from PIL import ImageDraw, ImageFont
    scale = 2
    W, H = CARD_W_PX * scale, CARD_H_PX * scale
    base = Image.open(BG_PATH).convert("RGB").resize((W, H), Image.LANCZOS)
    d = ImageDraw.Draw(base)

    def font(sz, bold=False):
        paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold
            else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for p in paths:
            try:
                return ImageFont.truetype(p, sz)
            except Exception:
                pass
        return ImageFont.load_default()

    blue = (0x04, 0x43, 0x6E)
    sky = (0x2E, 0x6E, 0xA8)
    grey = (0xB6, 0xC2, 0xCE)

    def ctext(y, txt, f, fill):
        w = d.textlength(txt, font=f)
        d.text(((W - w) / 2, y), txt, font=f, fill=fill)

    ctext(int(H * 0.10), "TRAINING FORMATION CENTER", font(40, True), blue)
    ctext(int(H * 0.205), "STUDENT IDENTITY CARD", font(26), sky)

    fields = ["NAME", "DATE OF BIRTH", "COURSE", "CARD ID", "VALID THRU", "SIGNATURE"]
    y = int(H * 0.31)
    lx = int(W * 0.10)
    for lab in fields:
        d.text((lx, y), lab, font=font(20, True), fill=blue)
        d.text((lx + int(W * 0.30), y), "______________________", font=font(22), fill=grey)
        y += int((H * 0.58) / len(fields))

    base.save("tfc_card_preview.png", dpi=(DPI, DPI))
    print("saved tfc_card_preview.png")


if __name__ == "__main__":
    build_background()
    build_doc()
    build_preview()
